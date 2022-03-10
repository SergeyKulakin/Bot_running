# -*- coding: utf-8 -*-
import telebot
from telebot import types
import sqlite3
import io #для работы с битовыми данными в памяти
import docx # для работы с Word
import datetime # для работы с временем
import stickers     # подключаем файл с стикерами
import notifications
import db_connection as db # подключаем db функции
from PIL import Image, ImageDraw, ImageFont # библиотека по работе с изображениями, добавление текста, шрифты


def db_initialize():
    conn = sqlite3.connect('running.db') # подключаем sqlite

    # курсор для работы с таблицами
    cursor = conn.cursor()

    running_data_query = "CREATE TABLE IF NOT EXISTS \"runningData\" (\"ID\" INTEGER UNIQUE, \"user_id\" INTEGER, \"run\" TEXT, PRIMARY KEY (\"ID\"))"
    reward_data_query  = "CREATE TABLE IF NOT EXISTS \"rewardData\" (\"ID\" INTEGER UNIQUE, \"user_id\" INTEGER, \"reward\" TEXT, PRIMARY KEY (\"ID\"))"

    # sql запрос на создание таблиц
    query_running_db = running_data_query
    cursor.execute(query_running_db)

    query_reward_db = reward_data_query
    cursor.execute(query_reward_db)


db_initialize()

# подключим токен нашего бота
bot = telebot.TeleBot(db.token)

# напишем, что делать нашему боту при команде старт
@bot.message_handler(commands=['start'])
def send_keyboard(message, text=notifications.bot_greeting):
    keyboard = types.ReplyKeyboardMarkup(row_width=1)
    lst_of_btns = ['Добавить пробежку', 'Удалить пробежку', 'Показать все пробежки', 'Показать пробежки за период',
                   'Показать активность за месяц', 'Импорт записей', 'Награды', 'Пока все!', 'Очистить runningData',
                   'Очистить rewardData']

    for btn_name in lst_of_btns:
        new_button = types.KeyboardButton(btn_name)
        keyboard.add(new_button)

    # пришлем это все сообщением и запишем выбранный вариант
    msg = bot.send_message(message.from_user.id, text=text, reply_markup=keyboard)

    # отправим этот вариант в функцию, которая его обработает
    bot.register_next_step_handler(msg, callback_worker)



#--------------------------------
# Теперь напишем функции для нашего бота

def check_marathon_flag(dist):
    return dist==42


def check_half_marathon_flag(dist):
    return dist==21

# функция "Добавить пробежку"
def add_run(msg):

    if (msg.text != 'Назад'):
        # выделим повторяющийся фрагмент в отдельную функцию
        flag_else, flag_except, flag_marathon, flag_half_marathon = add_run_repit_part(msg, msg.text)
        if flag_except and flag_else:
            bot.send_message(msg.chat.id, 'Вау, это надо зафиксировать :)')
            reward(msg, flag_marathon, flag_half_marathon)  # вызываем функцию учета наград
            #-----
            run_text = msg.text # сохраним текущую пробежку в переменную, потом передадим ее в функцию memoryPhoto
            markup = types.ReplyKeyboardMarkup()
            markup.add(types.KeyboardButton('Да'))
            markup.add(types.KeyboardButton('Нет'))
            msg = bot.send_message(msg.from_user.id, text='Сделаем памятную фотку?', reply_markup=markup)
            bot.register_next_step_handler(msg, memory_photo, run_text)  # переходим к функции memory photo

        elif flag_else == False:
            bot.send_message(msg.chat.id,
                             "Ваша дистанция - отрицательное число. Извините, я не могу добавить эту пробежку в список.")
            send_keyboard(msg, "Чем еще могу помочь?")
        elif flag_except == False:
            bot.send_message(msg.chat.id,
                             'Вы ввели пробежку в неправильном формате! Извините, я не могу её добавить в список')
            send_keyboard(msg, "Чем еще могу помочь?")

    else:
        send_keyboard(msg, "Хорошо, отменяем. Чем еще могу помочь?")

# вспомогательная функция для add_run
def add_run_repit_part(msg, text):
    flag_marathon = 0      # переменная, которая нужна для отслеживания, пробежал ли пользователь марафон
    flag_half_marathon = 0 # переменная, которая нужна для отслеживания, пробежал ли пользователь полумарафон
    flag_except = True
    flag_else = True
    try:
        # пробежка поступает в формате строки: "12 км, 44 минуты, Москва, 12.11.2020"
        # необходимо достать оттуда числовое значение дистанции пробежки
        # сначала достанем дистанцию в формате "12 км"
        run_dist_in_str_format = str(msg.text).split(',')[0]

        # теперь надо достать оттуда число
        digit_in_str_format = run_dist_in_str_format.split()[0]

        # переведем дистанцию в числовой формат
        cur_dist = float(digit_in_str_format)
        if cur_dist >= 0:
            if ',' in str(text).rstrip(')').rstrip(','):
                if cur_dist >= 42:
                    flag_marathon = 1
                elif 21 <= cur_dist < 42:
                    flag_half_marathon = 1
            db.insert_run_into_db(msg, text)

        else:
            flag_else = False
    except:
        flag_except = False

    return flag_else, flag_except, flag_marathon, flag_half_marathon

# функция для выбора сделать фото или нет
def memory_photo(msg, run_text):
    # bot.send_message(msg.chat.id, 'Попали в функцию memoryPhoto')
    if (msg.text == 'Да'):
        bot.send_message(msg.chat.id, 'Хорошо, скинь фотку... ')
        bot.register_next_step_handler(msg, image_open, run_text)
        #send_keyboard(msg, 'Могу помочь чем-то еще?')
    else:
        bot.send_message(msg.chat.id, 'Нет так нет')
        send_keyboard(msg, 'Могу помочь чем-то еще?')

# реализуем функцию "Обработки изображений"
@bot.message_handler(content_types=['photo'])
def image_open(msg, run_text):
    if msg.text != 'Назад':
        try:
            file_info = bot.get_file(msg.photo[len(msg.photo) - 1].file_id)
            downloaded_file = bot.download_file(file_info.file_path) #сохраняем фото с сервера телеграмм

            img = add_text_to_foto(run_text, downloaded_file)
            bot.send_photo(msg.chat.id, img)
            img.close()  # закрываем файл
            bot.send_message(msg.chat.id, "Ваше фото готово!")
            send_keyboard(msg, 'Могу помочь чем-то еще?')
        except Exception as e:
            bot.reply_to(msg, e)

    else:
        bot.send_message(msg, "Хорошо, отменяем. Чем еще могу помочь?")

# вспомогательная функция для image_open
def add_text_to_foto(run_text, downloaded_file):
    fp = io.BytesIO(downloaded_file)  # декодируем изображение из байт
    img = Image.open(fp)  # откроем изображение с библиотекой Pillow
    font = ImageFont.load_default()  # стандартный шрифт
    text = str(run_text)[-10:].encode('cp1251')  # .encode('UTF-8') #добавляемый шрифт
    draw_text = ImageDraw.Draw(img)  # добавим текст на изображение
    draw_text.text((10, 10), text, font=font, fill=('#58F9F6'))  # параметры размещения текста на изображении

    return img

"""Функция "Импорт из файла"
    напишем функцию "Импорт записей" для загрузки записей из файлов txt и docx.
    Это функция красивых строк для текстовых файлов(docx)"""
def get_document_string(documentList):
    document_list_str1 = [line.rstrip() for line in documentList]
    document_list_str2 = [str(val) for val in
                          document_list_str1]
    return '\n'.join(document_list_str2)

# Это функция красивых строк для текстовых файлов(txt)
def get_document_string_txt(document_list):
    document_list_str = [str(val) for val in
                         document_list]
    return '\n'.join(document_list_str)

# функция записи пробежек из файла
@bot.message_handler(content_types=['document'])
def import_run(msg):
    text = []
    if (msg.text != 'Назад'):
        bot.send_message(msg.chat.id, "Вы загрузили документ")
        try:
            chat_id = msg.chat.id
            file_info = bot.get_file(msg.document.file_id) #получаем информацию о файле на сервере telegram
            downloaded_file = bot.download_file(file_info.file_path) #записываем в память битовое представление файла
            end_file_txt = msg.document.file_name.endswith('.txt') #расширение файла txt
            end_file_docx = msg.document.file_name.endswith('.docx') #расширение файла docx
            if (end_file_txt == True) or (end_file_docx == True):

                if end_file_txt == True:
                    text = downloaded_file.decode('utf-8').split('\n')
                    doc_str = get_document_string_txt(text)
                    bot.send_message(msg.chat.id, 'Я сохраню эти записи:')
                    bot.send_message(msg.chat.id, doc_str)

                elif end_file_docx == True:
                     doc = docx.Document(io.BytesIO(downloaded_file))
                     text = [paragraph.text for paragraph in doc.paragraphs]
                     doc_str = get_document_string(text)
                     bot.send_message(msg.chat.id, 'Я сохраню эти записи:')
                     bot.send_message(msg.chat.id, doc_str)

            else: #если расширение другое, бот отправит стикер sorry
                bot.send_sticker(msg.chat.id, stickers.sorry_sticker)
        except Exception as e:
            return send_keyboard(msg, "Не могу прочесть. Чем еще могу помочь?")

        import_add_to_db(msg, text)

    else:
        send_keyboard(msg, "Хорошо, отменяем. Чем еще могу помочь?")

# вспомогательная функция import_run для зхаписи из файла в БД
def import_add_to_db(msg, text):
    record_run = [line.rstrip() for line in text] #запишем построчно данные из файла в БД
    flag_else, flag_except, flag_marathon, flag_half_marathon = True, True, 0, 0
    for rec in record_run:
        flag_else, flag_except, flag_marathon, flag_half_marathon = add_run_repit_part(msg, rec)
        reward(msg, flag_marathon, flag_half_marathon)

    if flag_else or flag_except:
        bot.send_message(msg.chat.id, 'Я всё зафиксировал :)')
        send_keyboard(msg, 'Могу помочь чем-то еще?')

    elif flag_else == False:
        bot.send_message(msg.chat.id,
                         "Дистанция в записях- отрицательное число. Извините, я не могу добавить эту пробежку в список.")
        send_keyboard(msg, "Чем еще могу помочь?")

    elif flag_except == False:
        bot.send_message(msg.chat.id,
                         'Есть пробежки в неправильном формате! Извините,'
                         'их я не могу добавить в список')
        send_keyboard(msg, "Чем еще могу помочь?")


# реализуем теперь функцию "Показать все пробежки"
# для этого сначала напишем функцию get_runs_string, которая делает красивые строки и отправляет их пользователю
def get_runs_string(runs_list):
    runs_list_str = []
    for val in list(enumerate(runs_list)):
        runs_list_str.append(str(val[0] + 1) + ') ' + val[1][0] + '\n')
    return ''.join(runs_list_str)

# теперь отправляем пользователю список пробежек
def show_all_runs(msg):
    runs_list = get_runs_string(db.select_runs_from_db(msg))
    bot.send_message(msg.chat.id, runs_list)
    send_keyboard(msg, "Чем еще могу помочь?")


# реализуем функцию "Удалить пробежку"
# для этого сначала пользователь должен выбрать пробежку, которую он хочет удалить
def choose_to_delete(msg):
    markup = types.ReplyKeyboardMarkup()
    runs_list = db.select_runs_from_db(msg)
    if len(runs_list) != 0:
        for val in runs_list:
            markup.add(types.KeyboardButton(val[0]))
        msg = bot.send_message(msg.from_user.id, text="Выбери пробежку для удаления", reply_markup=markup)
        bot.register_next_step_handler(msg, confirm_delete)
    else:
        bot.send_message(msg.from_user.id,'Здесь пусто. Нечего удалить.')
        send_keyboard(msg, "Чем еще могу помочь?")

# спрашиваем, точно ли пользователь хочет удалить пробежку
def confirm_delete(msg):
    bot.send_message(msg.from_user.id, 'Вы собираетесь удалить следующую пробежку:')
    bot.send_message(msg.from_user.id, msg.text)
    markup = types.ReplyKeyboardMarkup()
    markup.add(types.KeyboardButton('Да'))
    markup.add(types.KeyboardButton('Нет'))
    msg1 = bot.send_message(msg.from_user.id, text='Удаляем?', reply_markup=markup)
    bot.register_next_step_handler(msg1, delete_run, msg)

# теперь напишем саму функцию deleteRun
def delete_run(msg1, msg):
    if msg1.text == 'Да':
        db.delete_run_from_db(msg)
        bot.send_message(msg.chat.id, 'Пробежка удалена')
        send_keyboard(msg, "Чем еще могу помочь?")
    else:
        send_keyboard(msg, "Окей, ничего не удаляем! Чем еще могу помочь?")


# перейдем к функции Получения, просмотра и учета наград
# сначала добавим функцию для красивого вывода строк с наградой
def get_reward_string(reward_list):
    rew_list_str = []
    for val in list(reward_list):
        rew_list_str.append(str('"' + val[0] + '"' + '\n'))
    return ''.join(rew_list_str)

# теперь переходим к самой функции
def reward(msg, flag_marathon, flag_half_marathon):
    runs_list = db.select_runs_from_db(msg)
    reward_list = get_reward_string(db.select_rewards_from_db(msg))

    summdistance = 0  # переменная, в которую будем записывать суммарное расстояние, которое уже пробежал пользователь
    for val in runs_list:
        # try:
        if ',' in str(val).rstrip(')').rstrip(','):
            summdistance += float(str(val).split(',')[0].split()[0])#.lstrip('(').lstrip("'"))
        else:
            bot.send_message(msg.chat.id, notifications.wrong_format)
            bot.send_message(msg.chat.id, val)
            bot.send_message(msg.chat.id, "================")


    if (flag_marathon == 1) and ("Пробежать марафон!" not in reward_list):
        bot.send_message(msg.chat.id, 'Да вы марафонец!')
        bot.send_sticker(msg.chat.id, stickers.marathon_sticker)
        db.insert_reward_into_db(msg, "Пробежать марафон!")

    if (flag_half_marathon == 1) and ("Пробежать полумарафон!" not in reward_list):
        bot.send_message(msg.chat.id, 'Ого, вы пробежали полумарафон!')
        bot.send_sticker(msg.chat.id, stickers.half_marathon_sticker)
        db.insert_reward_into_db(msg, "Пробежать полумарафон!")

    if (summdistance > 0 and "Начало положено!" not in reward_list):
        bot.send_message(msg.chat.id, 'Ваша первая пробежка, круто!')
        bot.send_sticker(msg.chat.id, stickers.beginning_sticker)
        db.insert_reward_into_db(msg, "Начало положено!")

    if (summdistance >= 100 and "Преодолеть отметку в 100 км!" not in reward_list):
        bot.send_message(msg.chat.id, f'Вау, вы преодолели отметку в 100 км! Так держать!')
        bot.send_sticker(msg.chat.id, stickers.more_than_100_sticker)
        db.insert_reward_into_db(msg, "Преодолеть отметку в 100 км!")

    if (summdistance >= 200 and "Преодолеть отметку в 200 км!" not in reward_list):
        bot.send_message(msg.chat.id, 'Вау, вы преодолели отметку в 200 км! Очень круто!')
        bot.send_sticker(msg.chat.id, stickers.more_than_200_sticker)
        db.insert_reward_into_db(msg, "Преодолеть отметку в 200 км!")

# функция, которая выводит список всех полученных наград
def show_rewards(msg):
    reward_list = get_reward_string(db.select_rewards_from_db(msg))
    bot.send_message(msg.chat.id, reward_list)
    send_keyboard(msg, "Чем еще могу помочь?")


# напишем теперь функцию, которая удаляет ВСЕ пробежки из БД
# для этого сначала добавим функцию, которая будет спрашивать, уверен ли пользователь,
# что он хочет удалить все пробежки
def confirm_delete_all_running_data(msg):
    markup = types.ReplyKeyboardMarkup()
    markup.add(types.KeyboardButton('Да'))
    markup.add(types.KeyboardButton('Нет'))
    msg = bot.send_message(msg.from_user.id, text='Вы уверены, что хотите удалить ВСЕ пробежки?', reply_markup=markup)
    bot.register_next_step_handler(msg, delete_all_running_data)

# сама функция удаления пробежек
def delete_all_running_data(msg):
    if (msg.text == 'Да'):
        db.delete_all_runs_from_db(msg)
        send_keyboard(msg, "Все пробежки удалены. Чем еще могу помочь?")
    else:
        send_keyboard(msg, "Ок, ничего не удаляем. Чем еще могу помочь?")


# напишем теперь функцию, которая удаляет ВСЕ награды из БД
# для этого сначала добавим функцию, которая будет спрашивать, уверен ли пользователь,
# что он хочет удалить все награды
def confirm_delete_all_reward_data(msg):
    markup = types.ReplyKeyboardMarkup()
    markup.add(types.KeyboardButton('Да'))
    markup.add(types.KeyboardButton('Нет'))
    msg = bot.send_message(msg.from_user.id, text='Вы уверены, что хотите удалить ВСЕ достижения?', reply_markup=markup)
    bot.register_next_step_handler(msg, delete_all_reward_data)

# сама функция удаления наград
def delete_all_reward_data(msg):
    if (msg.text == 'Да'):
        db.delete_all_rewards_from_db(msg)
        send_keyboard(msg, "Все награды удалены. Чем еще могу помочь?")
    else:
        send_keyboard(msg, "Ок, ничего не удаляем. Чем еще могу помочь?")


# функция создания сводки
def activity_while_month(msg):
    runs_list = db.select_runs_from_db(msg)
    distance = 0
    time = 0
    places = []
    for i in runs_list:
        for j in i:
            try:
                info = str(j).split(',')
                date = info[3].replace(' ','').split('.')
                now = datetime.datetime.now().strftime('%y-%m-%d')
                bd = datetime.date(int(date[2]),int(date[1]),int(date[0])).strftime('%y-%m-%d')
                if (datetime.datetime.strptime(now,'%y-%m-%d') - datetime.datetime.strptime(bd,'%y-%m-%d')).days <= 30:
                    distance += float(str(i).split(',')[0].split()[0].lstrip('(').lstrip("'"))
                    time += int(str(i).split(',')[1].split()[0].lstrip('(').lstrip("'"))
                    places.append(str(info[2]))
            except:
                # bot.send_message(msg.chat.id, "Следующая пробежка не удовлетворяет требованиям к формату, поэтому не будет учтена")
                # bot.send_message(msg.chat.id, i)
                pass

    if len(places) != 0:
        max_place = max(set(places), key=places.count)
    else:
        max_place = 'Нет мест'
    bot.send_message(msg.chat.id, 'Расстояние: {} км'.format(distance))
    bot.send_message(msg.chat.id, 'Время: {} минут'.format(time))
    bot.send_message(msg.chat.id, 'Самое часто посещаемое место: {}'.format(max_place))
    send_keyboard(msg, "Чем еще могу помочь?")

# програмка на вход получает количество месяцев, а дальше смотрит на разницу по датам и подбирает все пробежки
def runs_in_period(msg):
    answer = ''
    runs_list = db.select_runs_from_db(msg)
    if len(runs_list) != 0:
        spisok = []
        for i in runs_list:
            for j in i:
                try:
                    info = str(j).split(',')
                    date = info[3].replace(' ', '').split('.')
                    now = datetime.datetime.now().strftime('%y-%m-%d')
                    bd = datetime.date(int(date[2]), int(date[1]), int(date[0])).strftime('%y-%m-%d')
                    if (datetime.datetime.strptime(now, '%y-%m-%d') - datetime.datetime.strptime(bd, '%y-%m-%d')).days <= int(msg.text) * 30:
                        spisok.append(j)
                except:
                    pass

        if len(spisok) != 0:
            for i in spisok:
                answer += str(i) + '\n'
            bot.send_message(msg.from_user.id, answer)
            send_keyboard(msg, "Чем еще могу помочь?")
        else:
            bot.send_message(msg.from_user.id, "У вас нет пробежек за указанный период")
            send_keyboard(msg, "Чем еще могу помочь?")
    else:
        bot.send_message(msg.from_user.id, "У вас еще нет пробежек")
        send_keyboard(msg, "Чем еще могу помочь?")
#--------------------------------

# привязываем функции к кнопкам на клавиатуре
def callback_worker(call):
    if call.text == "Добавить пробежку":
        markup = types.ReplyKeyboardMarkup(row_width=1)
        markup.add(types.KeyboardButton('Назад'))
        msg = bot.send_message(call.chat.id, text=notifications.add_run, reply_markup=markup)
        bot.register_next_step_handler(msg, add_run)

    # кнопка импорта файлов
    elif call.text == "Импорт записей":
        markup = types.ReplyKeyboardMarkup(row_width=1)
        markup.add(types.KeyboardButton('Назад'))
        msg = bot.send_message(call.chat.id, notifications.import_runs, reply_markup=markup)
        bot.register_next_step_handler(msg, import_run)


    elif call.text == "Показать все пробежки":
        try:
            show_all_runs(call)
        except:
            bot.send_message(call.chat.id, 'Здесь пусто. Может начнем бежать прямо сейчас? ;)')
            send_keyboard(call, "Чем еще могу помочь?")


    elif call.text == "Удалить пробежку":
        choose_to_delete(call)

    elif call.text == "Награды":
        bot.send_message(call.chat.id, 'Давайте посмотрим на ваши награды!...')
        try:
            show_rewards(call)
        except:
            bot.send_message(call.chat.id, 'У вас пока нет наград :(')
            bot.send_sticker(call.chat.id, stickers.sad_sticker)
            send_keyboard(call, "Чем еще могу помочь?")

    elif call.text == 'Показать активность за месяц':
        bot.send_message(call.chat.id, 'Вот ваша активность за последний месяц')
        activity_while_month(call)


    elif call.text == 'Показать пробежки за период':
        msg = bot.send_message(call.from_user.id, 'Введите количество месяцев')
        bot.register_next_step_handler(msg, runs_in_period)


    elif call.text == "Пока все!":
        bot.send_message(call.chat.id, 'Хорошего дня! Когда захотите продолжить нажмите на команду\n /start')

    elif call.text == 'Очистить runningData':
        confirm_delete_all_running_data(call)

    elif call.text == 'Очистить rewardData':
        confirm_delete_all_reward_data(call)

    else:
        bot.send_message(call.chat.id, 'Извините, я вас не понимаю :(')
        send_keyboard(call, "Пожалуйста, выберите функции из предложенных")


bot.polling(none_stop=True)
